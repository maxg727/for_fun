import streamlit as st
import sqlite3
from pathlib import Path
import sys

from modules import rulebook
from modules.weekly_matchup import WeeklyMatchupModule

# Add modules directory to path
sys.path.append('modules')

# Import dashboard modules
from modules.standings import StandingsModule
from modules.team_analysis import TeamAnalysisModule


# from matchups import MatchupsModule  # Future module
# from trades import TradesModule      # Future module
# from players import PlayersModule    # Future module

# Page configuration
st.set_page_config(
    page_title="Sleeper Fantasy Dashboard",
    page_icon="🏈",
    layout="wide",
    initial_sidebar_state="expanded"
)


class SleeperDashboard:
    def __init__(self, db_path="sleeper_data.db"):
        self.db_path = db_path
        self.check_database()
        self.init_session_state()

    def check_database(self):
        """Check if database exists and has data"""
        if not Path(self.db_path).exists():
            st.error(f"""
            🚨 Database not found: `{self.db_path}`

            Please run the data pipeline first:
            ```bash
            python sleeper_pipeline.py
            ```
            """)
            st.stop()

        # Check if database has data
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM leagues")
        league_count = cursor.fetchone()[0]
        conn.close()

        if league_count == 0:
            st.error("""
            🚨 Database is empty!

            Please run the data pipeline to populate data:
            ```bash
            python sleeper_pipeline.py
            ```
            """)
            st.stop()

    def init_session_state(self):
        """Initialize session state variables"""
        if 'selected_league' not in st.session_state:
            st.session_state.selected_league = None
        if 'selected_season' not in st.session_state:
            st.session_state.selected_season = None

    def get_available_leagues(self):
        """Get all available leagues from database"""
        conn = sqlite3.connect(self.db_path)
        query = """
            SELECT DISTINCT league_id, name, season 
            FROM leagues 
            ORDER BY season DESC, name
        """
        df = pd.read_sql_query(query, conn)
        conn.close()
        return df

    def render_sidebar(self):
        """Render sidebar with league/season selection"""
        st.sidebar.title("🏈 Dashboard Controls")

        # Get available leagues
        leagues_df = self.get_available_leagues()

        if leagues_df.empty:
            st.sidebar.error("No leagues found in database")
            return None, None

        # Season selector
        seasons = sorted(leagues_df['season'].unique(), reverse=True)
        selected_season = st.sidebar.selectbox(
            "📅 Select Season",
            seasons,
            key="season_selector"
        )

        # League selector (filtered by season)
        season_leagues = leagues_df[leagues_df['season'] == selected_season]
        league_options = {}
        for _, row in season_leagues.iterrows():
            league_options[f"{row['name']} ({row['season']})"] = row['league_id']

        if league_options:
            selected_league_name = st.sidebar.selectbox(
                "🏆 Select League",
                list(league_options.keys()),
                key="league_selector"
            )
            selected_league_id = league_options[selected_league_name]
        else:
            st.sidebar.error(f"No leagues found for {selected_season}")
            return None, None

        # Update session state
        st.session_state.selected_league = selected_league_id
        st.session_state.selected_season = selected_season

        # Database info
        with st.sidebar.expander("📊 Database Info"):
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()

            # Get table counts
            tables = ['leagues', 'rosters', 'matchups', 'transactions', 'players']
            for table in tables:
                cursor.execute(f"SELECT COUNT(*) FROM {table}")
                count = cursor.fetchone()[0]
                st.info(f"**{table.title()}**: {count:,}")

            # Last refresh info
            cursor.execute("""
                SELECT completed_at FROM refresh_log 
                WHERE success = 1 
                ORDER BY completed_at DESC LIMIT 1
            """)
            result = cursor.fetchone()
            if result:
                st.success(f"**Last Updated**: {result[0]}")

            conn.close()

        # Refresh button
        if st.sidebar.button("🔄 Refresh Data", help="Run data pipeline to update"):
            st.sidebar.info("Run: `python sleeper_pipeline.py`")

        return selected_league_id, selected_season

    def render_header(self, league_id, season):
        """Render dashboard header with league info"""
        conn = sqlite3.connect(self.db_path)
        query = """
            SELECT name, total_rosters, scoring_settings, status
            FROM leagues 
            WHERE league_id = ?
        """
        cursor = conn.cursor()
        cursor.execute(query, (league_id,))
        result = cursor.fetchone()
        conn.close()

        if result:
            name, total_rosters, scoring_settings, status = result

            # Parse scoring type
            try:
                import json
                scoring = json.loads(scoring_settings) if scoring_settings else {}
                rec_points = scoring.get('rec', 0)
                if rec_points == 1:
                    scoring_type = "PPR"
                elif rec_points == 0.5:
                    scoring_type = "Half PPR"
                else:
                    scoring_type = "Standard"
            except:
                scoring_type = "Unknown"

            # Header with league info
            col1, col2, col3, col4 = st.columns([3, 1, 1, 1])

            with col1:
                st.title(f"🏈 {name}")
                st.caption(f"Season {season}")

            with col2:
                st.metric("Teams", total_rosters or "N/A")

            with col3:
                st.metric("Scoring", scoring_type)

            with col4:
                status_color = "🟢" if status == "in_season" else "🟡"
                st.metric("Status", f"{status_color} {status.replace('_', ' ').title()}")

    def run(self):
        """Main dashboard execution"""
        # Render sidebar and get selections
        league_id, season = self.render_sidebar()

        if not league_id or not season:
            st.warning("Please select a league to view dashboard")
            return

        # Render header
        self.render_header(league_id, season)

        tab1, tab2, tab3, tab4 = st.tabs([
            "📊 Standings",
            "📈 Team Analysis",
            "⚡ Matchups",
            "📚Constitution"
        ])

        # Standings
        with tab1:
            standings_module = StandingsModule(self.db_path)
            standings_module.render(league_id, season)

        # Team Analysis
        with tab2:
            team_analysis_module = TeamAnalysisModule(self.db_path)
            team_analysis_module.render(league_id, season)

        with tab3:
            weekly_matchup_module = WeeklyMatchupModule(self.db_path)
            weekly_matchup_module.render(league_id, season)

        # Rulebook
        with tab4:
            import modules.rulebook as rulebook
            st.subheader("League Constitution & Rulebook")
            st.markdown(rulebook.RULEBOOK, unsafe_allow_html=True)

            st.markdown("### 📥 Download Rulebook")
            col1, col2 = st.columns(2)

            # DOCX export
            with col1:
                from io import BytesIO
                try:
                    from docx import Document
                except ImportError:
                    st.error("python-docx not installed. Run `pip install python-docx`.")
                else:
                    doc = Document()
                    doc.add_heading("League Constitution & Rulebook", 0)
                    doc.add_paragraph(rulebook.RULEBOOK)

                    docx_buffer = BytesIO()
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)

                    st.download_button(
                        label="📄 Download as DOCX",
                        data=docx_buffer,
                        file_name="league_rulebook.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            # PDF export
            with col2:
                try:
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet
                except ImportError:
                    st.error("reportlab not installed. Run `pip install reportlab`.")
                else:
                    pdf_buffer = BytesIO()
                    pdf = SimpleDocTemplate(pdf_buffer)
                    styles = getSampleStyleSheet()
                    story = [Paragraph("League Constitution & Rulebook", styles['Title']), Spacer(1, 12)]

                    for line in rulebook.RULEBOOK.split("\n"):
                        if line.strip():
                            story.append(Paragraph(line, styles['Normal']))
                            story.append(Spacer(1, 6))

                    pdf.build(story)
                    pdf_buffer.seek(0)

                    st.download_button(
                        label="📕 Download as PDF",
                        data=pdf_buffer,
                        file_name="league_rulebook.pdf",
                        mime="application/pdf"
                    )



def main():
    """Main entry point"""
    # Import pandas here to avoid issues if not installed
    try:
        import pandas as pd
        globals()['pd'] = pd
    except ImportError:
        st.error("Please install pandas: `pip install pandas`")
        st.stop()

    # Initialize and run dashboard
    dashboard = SleeperDashboard()
    dashboard.run()


if __name__ == "__main__":
    main()